#!/usr/bin/env python

#################################################################################################################
################---ASPECTOS POR MEJORAR---#######################################################################
#	Cantidad total de amino acidos 
#	Cantidad amino acidos por bloques (1-10)
#	Informacion de cada acople
#################################################################################################################
import pandas as pd
from collections import Counter
from collections import defaultdict
from Bio.SeqUtils import molecular_weight
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.shared import qn
from io import BytesIO

# Diccionario de aminoácidos
AA_DICT = {
    "A": "ALA", "R": "ARG", "N": "ASN", "D": "ASP", "C": "CYS",
    "Q": "GLN", "E": "GLU", "G": "GLY", "H": "HIS", "I": "ILE",
    "L": "LEU", "K": "LYS", "M": "MET", "F": "PHE", "P": "PRO",
    "S": "SER", "T": "THR", "W": "TRP", "Y": "TYR", "V": "VAL"
}
##	El peso molecular de los AA están con los grupos protectores Fmoc y :
##	R=R(Pbf), N=N(Trt), D=D(tBu), C=C(Trt), E=E(tBu), Q=Q(Trt), H=H(Trt)
##	K=K(Boc), S=S(tBu), Y=Y(tBu), T=T(tBu), W=W(Boc)
AA_MW = {
    "A": 311.3, "R": 648.8, "N": 596.7, "D": 411.5,
    "C": 585.7, "E": 443.5, "Q": 610.7, "G": 297.3,
    "H": 619.7, "I": 353.4, "L": 353.4, "K": 468.5,
    "M": 371.5, "F": 387.4, "P": 337.4, "S": 383.4,
    "T": 397.5, "W": 526.6, "Y": 459.6, "V": 339.4
}
###########################################################################################
#######------GENERATE INTRO-INFO-----#########
def intro(doc, nameProject, desprotection, nameResin, massResin, StResin, peptides):
    def add_line(label, value):
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(str(value))
    add_line("Nombre de Síntesis:\t\t\t\t", nameProject)
    add_line("Método·de·Desprotección:\t\t\t", desprotection)
    add_line("Nombre de Resina:\t\t\t\t", nameResin)
    add_line("Cantidad de Resina (mg):\t\t\t", float(massResin))
    add_line("Sustitución de Resina (mmol/g):\t\t", StResin)
    add_line("Cantidad de péptidos:\t\t\t", len(peptides))

###########################################################################################
###############----GENERATE TABLE OF PEPTIDES----#########
def introtable(doc, bolsas, peptides, family, notes, StResin, massResin):
    ##==== Titulo antes de la tabla
    title_table=doc.add_heading("PÉPTIDOS A SINTETIZAR", level=1)
    title_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ##==== Crear tabla
    table = doc.add_table(rows=len(peptides)+1, cols=7)
    table.autofit = True	#False
    table.style = "Table Grid"
    #===== Encabezados
    headers = ["N° Bolsa", "Secuencia", "Largo", "MW (g/mol)", "Familia", "Masa teórica (mg)", "Nota"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Courier New' # Define el tipo de letra
        run.font.size = Pt(10)        # Define el tamaño
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


    #===== Listado de peptidos
    for i, (bolsa,seq,fam,nota) in enumerate(zip(bolsas, peptides, family,notes)):
        row = table.rows[i+1].cells
        mw = round(molecular_weight(seq, seq_type="protein"), 2)
        mt = round(StResin * mw * massResin / 1000, 2)
        row[0].text = str(bolsa)        	# Bolsa
        #row[0].text = str(i+1) 		# Bolsa
        row[1].text = seq 			# Secuencia Peptido
        row[2].text = str(len(seq)) 		# Largo
        row[3].text = f"{mw:.2f}" 		# Peso molecular
        row[4].text = fam            		# Familia del peptido
        row[5].text = f"{mt:.2f}"  		# Masa Teorica ===REVISAR====
        row[6].text = str(nota) if nota else "" # Nota adicional por peptido

    #===== Formato global de la tabla, alineacion
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()
    for row in table.rows:
        row.height = Mm(4)
###########################################################################################
#######------TABLA DE TOTAL DE AMINOACIDOS-----#########
def count_total_aminoacids(peptides):
    total_counts = Counter()
    for pep in peptides:
        total_counts.update(pep)
    return total_counts
###########################################################################################
def calculate_mass_excess(counts, StResin, mass_resin, excess=10):
    result = {}
    for aa, count in counts.items():
        mw = AA_MW.get(aa, 0)
        mass = StResin * mass_resin * mw * excess * count / 1000
        result[aa] = round(mass, 2)
    return result
###########################################################################################
def add_total_aa_table(doc, peptides, StResin, mass_resin):
    # 1. Obtener conteos y cálculos de masa
    counts = count_total_aminoacids(peptides)
    mass_calc = calculate_mass_excess(counts, StResin, mass_resin)
    title = doc.add_heading("CONSUMO TOTAL DE AMINOÁCIDOS", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 2. 🔥 ORDENAMIENTO ALFABÉTICO por nombre de 3 letras (ALA, ARG...)
    # Usamos AA_DICT para que el criterio de orden sea el nombre completo
    sorted_items = sorted(counts.items(), key=lambda item: AA_DICT.get(item[0], item[0]))
    # 3. Crear tabla de 3 columnas
    table = doc.add_table(rows=len(counts) + 1, cols=3)
    table.style = "Table Grid"
    ### Encabezado ###
    headers = ["Aminoácido", "N° Residuos", "Masa (10x) mg"]
    table.rows[0].height = Mm(10) # Un poco más alto para el título
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = True
    ### Contenido ###
    # 🔥 IMPORTANTE: Iteramos sobre sorted_items (la lista ya ordenada)
    for i, (aa, count) in enumerate(sorted_items):
        row = table.rows[i+1]
        row.height = Mm(5) # Grosor de fila para legibilidad
        cells = row.cells
        # Nombre de 3 letras (ALA, 
        cells[0].text = AA_DICT.get(aa, aa)
        cells[1].text = str(count)
        # Masa calculada (aseguramos 1 decimal)
        m_val = mass_calc.get(aa, 0.0)
        cells[2].text = f"{float(m_val):.1f}"
        # Alineación total de la fila
        for cell in cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)
    doc.add_page_break()
###########################################################################################
def count_aa_per_block(cycles, block_size=10):
    blocks = []
    # Creamos un mapa inverso de nombre -> letra (ej: 'Alanina': 'A')
    # Esto asume que tienes un diccionario inverso o lo creamos aquí
    INV_AA_DICT = {v: k for k, v in AA_DICT.items()}
    for i in range(0, len(cycles), block_size):
        block_cycles = cycles[i:i+block_size]
        block_counter = Counter()
        for aa_positions, _, _ in block_cycles:
            for aa_name, peps in aa_positions.items():
                # 🔥 PASO CLAVE: Convertir el nombre a letra para que coincida con FMOC_MW
                # Si aa_name ya es la letra, esto no hará nada malo
                aa_letter = INV_AA_DICT.get(aa_name, aa_name)
                block_counter[aa_letter] += len(peps)
        blocks.append((i+1, i+len(block_cycles), block_counter))
    return blocks

###########################################################################################
def add_block_tables(doc, cycles, StResin, mass_resin):
    # 1. Obtener los bloques (devuelve conteos por letra: 'A', 'R', etc.)
    blocks = count_aa_per_block(cycles, block_size=10)
    for start, end, counts in blocks:
        title = doc.add_heading(f"CANTIDAD DE AMINOÁCIDOS: ACOPLE {start} AL {end}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 2. Calcular masa (usa letras para FMOC_MW)
        mass_calc = calculate_mass_excess(counts, StResin, mass_resin)
        # 3. Crear tabla
        table = doc.add_table(rows=len(counts) + 1, cols=3)
        table.style = "Table Grid"
        # --- Formato Encabezado ---
        headers = ["Aminoácido", "Residuos", "Masa (10x) mg"]
        table.rows[0].height = Mm(10)
        # --- Encabezado ---
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
        # --- Contenido de la Tabla (ORDENADO ALFABÉTICAMENTE) ---
        # 🔥 CAMBIO AQUÍ: La 'key' del sorted usa AA_DICT para que 'ALA' vaya antes que 'ARG'
        sorted_block = sorted(
            counts.items(), 
            key=lambda item: AA_DICT.get(item[0], item[0])
        )
        for i, (aa_letter, count) in enumerate(sorted_block):
            row = table.rows[i+1]
            row.height = Mm(5) # Grosor para legibilidad
            cells = row.cells
            # Nombre de 3 letras para la tabla (ALA, ARG...)
            nombre_visual = AA_DICT.get(aa_letter, aa_letter)
            cells[0].text = nombre_visual 
            cells[1].text = str(count)
            # Masa calculada con la letra ('A')
            m_val = mass_calc.get(aa_letter, 0.0)
            cells[2].text = f"{float(m_val):.1f}"
            # Centrado total
            for cell in cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()
#####################################################################################
#######------GENERATE COUPLINGS-----#########
###########################################################################################
def get_cycles(peptides, bolsas):
    reversed_peptides = [p[::-1] for p in peptides]
    max_length = max(len(p) for p in peptides)
    cycles = []
    for i in range(max_length):
        aa_positions = defaultdict(list)
        ending_peptides = []
        active_peptides = []
        for idx, pep in enumerate(reversed_peptides):
            if i < len(pep):
                aa = pep[i]
                bolsa = bolsas[idx]
                aa_positions[AA_DICT[aa]].append(bolsa)
                active_peptides.append(bolsa)
                if i == len(pep) - 1:
                    ending_peptides.append(bolsa)
        cycles.append((aa_positions, ending_peptides, active_peptides))
    return cycles

###########################################################################################
def write_couplings(doc, cycles, bolsas, peptides, simple, doble, triple, desprotection):

    # 🔥 CONFIGURACIÓN LOCAL DE ESTA FUNCIÓN
    font_name = "Cambria"
    font_size = 10
    line_spacing = 1.0

    def format_cell(cell):
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.line_spacing = line_spacing
            paragraph.paragraph_format.space_before = Pt(1)
            paragraph.paragraph_format.space_after = Pt(1)

            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)

    # 🔥 ANTES DEL LOOP
    active_bags = set(bolsas)

    for i, (aa_positions, ending, active) in enumerate(cycles):

        tot_bolsas = len(active_bags)

        title = doc.add_heading(
            f"ACOPLE NÚMERO {i+1} \t\t\t{tot_bolsas} bolsas",
            level=1
        )

        # === OUT ===
        if ending:
            p_out = doc.add_paragraph()
            tab_stops = p_out.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Mm(30))

            out_str = "  ".join(map(str, sorted(ending)))
            run = p_out.add_run(f"OUT:\t{out_str}")
            run.bold = True
            run.font.size = Pt(8)

        # === HEADER AA ===
        p_header = doc.add_paragraph()
        tab_stops = p_header.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Mm(30))

        run_h = p_header.add_run("Aminoácido\tBolsas")
        run_h.bold = True
        p_header.paragraph_format.space_after = Pt(2)

        # === FILTRADO ===
        aa_filtered = {}
        for aa, peps in aa_positions.items():
            valid_peps = [p for p in peps if p in active_bags]
            if valid_peps:
                aa_filtered[aa] = valid_peps

        sorted_items = sorted(
            aa_filtered.items(),
            key=lambda x: AA_DICT.get(x[0], x[0])
        )

        # === CONTENIDO AA ===
        for aa, peps in sorted_items:
            p_row = doc.add_paragraph()
            p_row.paragraph_format.tab_stops.add_tab_stop(Mm(30))
            p_row.paragraph_format.line_spacing = 1.0
            p_row.paragraph_format.space_after = Pt(0)

            nombre_3 = AA_DICT.get(aa, aa)
            run = p_row.add_run(f"{nombre_3}\t")
            run.bold = True
            run.font.size = Pt(10)

            bolsas_str = "  ".join(map(str, peps))
            run2 = p_row.add_run(f"{bolsas_str}   =   {len(peps)}")
            run2.font.size = Pt(10)

        # 🔥 actualizar bolsas activas
        active_bags -= set(ending)

        doc.add_paragraph()

        # =========================================================
        # 🔥 TABLA DESPROTECCIÓN
        # =========================================================
        table = doc.add_table(rows=1, cols=5)
        table.autofit = True

        widths = [Mm(75),Mm(20), Mm(28), Mm(28), Mm(35)]
        for j, w in enumerate(widths):
            table.columns[j].width = w

        # quitar bordes
        for row in table.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ('top', 'left', 'bottom', 'right'):
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                tcPr.append(tcBorders)

        # header
        row = table.rows[0].cells
        row[0].text = "DESPROTECCIÓN"
        row[1].text = "FECHA"
        row[2].text = "CHEQUEO"
        row[3].text = "HECHO POR"
        row[4].text = "REVISADO POR"

        for cell in row:
        #    format_cell(cell)
             for paragraph in cell.paragraphs:
                 for run in paragraph.runs:
                     run.bold = True   # 🔥 NEGRILLA
#    format_cell(cell)
        def add_row(text, check):
            row_cells = table.add_row().cells
            row_cells[0].text = text
            row_cells[1].text = "___/___/___"
            row_cells[2].text = check
            row_cells[3].text = "_________________"
            row_cells[4].text = "_________________"

            for cell in row_cells:
                format_cell(cell)

        # contenido
        add_row(f"{desprotection} (2x10') ♻️", "|_____|_____|")
        add_row("Lavado con DMF (3x1') ♻️", "|___|___|___|")
        add_row("Lavado con IPA (1x1')", "|___________|")
        add_row("Lavado con BPB 1% en DMF (1x2')", "|___________|")
        add_row("Lavado con DMF (2x1') ♻️", "|_____|_____|")
        add_row("Lavado con DCM (1x1')", "|___________|")

        # =========================================================
        # 🔥 TABLA ACOPLE
        # =========================================================
        table2 = doc.add_table(rows=1, cols=5)
        table2.autofit = False

        for j, w in enumerate(widths):
            table2.columns[j].width = w

        # quitar bordes
        for row in table2.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ('top', 'left', 'bottom', 'right'):
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                tcPr.append(tcBorders)

        # header
        row = table2.rows[0].cells
        row[0].text = "\nCICLO DE ACOPLE"
        row[1].text = "\nFECHA"
        row[2].text = "\nHORA"
        row[3].text = "\nHECHO POR"
        row[4].text = "\nREVISADO POR"

        for cell in row:
           # format_cell(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True   # 🔥 NEGRILLA

        def add_row2(text):
            row_cells = table2.add_row().cells
            row_cells[0].text = text
            row_cells[1].text = "___/___/___"
            row_cells[2].text = "|___:___||___:___|"
            row_cells[3].text = "__________________"
            row_cells[4].text = "_________________"
            for cell in row_cells:
                format_cell(cell)

        # contenido
        add_row2("Simple")
        row_cells = table2.add_row().cells
        row_cells[0].text = simple
        for cell in row_cells:
            format_cell(cell)

        add_row2("Doble ♻️")
        row_cells = table2.add_row().cells
        row_cells[0].text = doble
        for cell in row_cells:
            format_cell(cell)

        add_row2("Triple ♻️")
        row_cells = table2.add_row().cells
        row_cells[0].text = triple
        for cell in row_cells:
            format_cell(cell)
#        doc.add_page_break()

#*del        # === FINALIZACIÓN ===
#*del        if ending:
#*del            p = doc.add_paragraph()
#*del            run = p.add_run("Finalizan bolsas: " + ", ".join(map(str, ending)))
#*del            run.bold = True
#*del            run.underline = True
#*del
#        doc.add_page_break()

        ## === INFO FINALIZACIÓN
        #if ending:
        #    p = doc.add_paragraph()
        #    run = p.add_run(
        #        "Finalizan bolsas: " + ", ".join(map(str, ending)) )
        #    run.bold = True
        #    run.underline = True
        #=== Informacion con stop de espaciado (tabs)
        p_header = doc.add_paragraph()
        tab_stops = p_header.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Mm(145))  # Posición para Bolsa(s)
        tab_stops.add_tab_stop(Mm(165)) # Posición para Total (cerca del margen derecho)
        # Escribimos el encabezado en negrita
       # run_h = p_header.add_run("Aminoácido\tBolsas\tTotal")
        run_h = p_header.add_run("\nBolsas : Aminoácido")
        run_h.bold = True
        p_header.paragraph_format.space_after = Pt(1) # Espacio pequeño bajo el encabezado
        # === CONTENIDO ORDENADO (3 LETRAS Y ALFABÉTICO) ===
        # Ordenamos por el nombre de 3 letras (ALA, ARG...)
        sorted_items = sorted(aa_positions.items(), key=lambda x: AA_DICT.get(x[0], x[0]))
###########################################################################################
        lista_plana = []
        for aa, peps in sorted_items:
            nombre_3 = AA_DICT.get(aa, aa)
            for bolsa in peps:
                lista_plana.append((int(bolsa), nombre_3))
        
        # 2. Ordenamos por número de bolsa para que la secuencia sea lógica
        lista_plana.sort(key=lambda x: x[0])
        
        # 3. Calculamos cuántas filas necesitamos (de a 10 elementos por fila)
        num_elementos = len(lista_plana)
        num_filas = 10
        num_columnas = (num_elementos + num_filas - 1) // num_filas

        # 4. Creamos la tabla en el documento
        table = doc.add_table(rows=num_filas, cols=num_columnas)
        table.style = 'Normal Table' #'Table Grid' # Opcional: añade bordes para guiar la vista
        
        table.autofit = True
        

        for i, (num_bolsa, nombre_aa) in enumerate(lista_plana):
            columna = i // num_filas
            fila = i % num_filas
            # Accedemos a la celda específica
            cell = table.cell(fila, columna)
            # Escribimos el formato "1:PHE"
            p = cell.paragraphs[0]
            p.style = doc.styles['No Spacing']
    #        p.paragraph_format.line_spacing = 1.0
          # Esto elimina el margen invisible que Word pone arriba y abajo del texto
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run =p.add_run(f"{num_bolsa}:{nombre_aa}")
            run.font.name = 'Courier New'
            #run = p.runs[0]
            run.font.size = Pt(8)
        doc.add_page_break()
###########################################################################################
def add_page_numbers(doc):
    # Accedemos a la sección actual (por defecto la primera)
    for section in doc.sections:
        footer = section.footer
        # Si el footer tiene párrafos, usamos el primero; si no, creamos uno
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # --- Inicio del código mágico de Word para el campo "PAGE" ---
        run = p.add_run()
        # Crear elementos XML para el campo dinámico
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE" # Este es el comando que le dice a Word: "Pon el número de página"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        # Insertar los elementos en el 'run'
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
###########################################################################################
def create_word(nameProject, desprotection, nameResin, massResin, StResin, bolsas, peptides, family, notes, simple, doble, triple, fileName):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(5)
    section.left_margin = Mm(15)
    section.right_margin = Mm(10)

    # === ciclos
    cycles = get_cycles(peptides, bolsas)
    # === secciones
    intro(doc, nameProject, desprotection, nameResin, massResin, StResin, peptides)	# Funcion para escribir la informacion del proyecto.
    #introtable(doc, peptides, family)							# Funcion para escribir la tabla de la síntesis
    introtable(doc, bolsas, peptides, family, notes, StResin, massResin)
    add_total_aa_table(doc, peptides, StResin, massResin)				# Funcion para escribir la tabla de consumo total
    add_block_tables(doc, cycles, StResin, massResin)					# Funcion para escribir la tabla de consumo por bloques de 10
    write_couplings(doc, cycles, bolsas, peptides, simple, doble, triple, desprotection)			# Funcion para escribir el texto para cada acople
    add_page_numbers(doc)								# Funcion para agregar N° de Pagina
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
#    doc.save(fileName)
#    return fileName

#*****#***###############-------FIN DEL CODIGO----#########
